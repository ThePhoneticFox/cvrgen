import cfg
import docx, os



def displayfiles(component):
    x = 1
    cwd = os.getcwd()
    pointpath = os.path.join(cwd, 'components', component)
    for filename in os.listdir(pointpath):
        #print(filename)
        y = str(x)
        readdoc = docx.Document(pointpath + '\\' + filename)
        heading = readdoc.paragraphs[0].text
        body = readdoc.paragraphs[1].text
        cfg.headerlist.append(heading)
        cfg.bodylist.append(body)
        print(y + '. ' + heading)
        x = x + 1


