#Brendon Berndt 2018 - Don't be a jerk, credit where applicable (MIT Licence)

import cfg
from cfg import yourName
import os
import docx
import reader, assembler
#reader = reader.spec_from_file_location("reader.py", work + "\\modules")
#cfg.init()
work = os.getcwd()
#reader.getfiles()
#TODO Custom address line for personalisation
print('\n')
greeting = input('Please type the greeting you would like (include commas)> ')
job = input('Please type the job name > ')

print(cfg.chosen)
#TODO Add extra paragraphs for choice logic
#TODO Organise components into file
header = open('cvheader.txt')
header = header.read()
jobheader = header.replace('<JOB>', job)
footer = open('cvfooter.txt')


#TODO Create newline after point 3 (one line, not two)
cfg.coverletter.add_paragraph(greeting)
cfg.coverletter.add_paragraph()
cfg.coverletter.add_paragraph(jobheader)
reader.displayfiles('points')
assembler.assembler()
#cfg.coverletter.add_paragraph()
cfg.coverletter.add_paragraph()
cfg.coverletter.add_paragraph(footer)
cfg.coverletter.add_paragraph()
cfg.coverletter.add_paragraph('Yours Sincerley,')
cfg.coverletter.add_paragraph(yourName)

#TODO Formatting

cfg.coverletter.save(yourName + 'Cover Letter 2018.docx')
