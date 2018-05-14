import cfg
import os
import docx
import reader

def assembler():
    for x in range(0, 3):
        out = input('Pick a paragraph: ')
        cfg.chosen = int(out) - 1
        cfg.chooselist.append(cfg.chosen)
        #print(cfg.bodylist[cfg.chosen])

    #for x in range(0, 3):
        #print(cfg.chooselist[x])
        #print(cfg.bodylist[cfg.chooselist[x]])

#TODO Assemble paragraphs for letters
    #Paragraph 1


    header1 = cfg.coverletter.add_paragraph(cfg.headerlist[cfg.chooselist[0]])
    header1.add_run(' - ')
    header1.add_run(cfg.bodylist[cfg.chooselist[0]])
    header1.runs[0].font.bold = True

    header2 = cfg.coverletter.add_paragraph(cfg.headerlist[cfg.chooselist[1]])
    header2.add_run(' - ')
    header2.add_run(cfg.bodylist[cfg.chooselist[1]])
    header2.runs[0].font.bold = True

    header3 = cfg.coverletter.add_paragraph(cfg.headerlist[cfg.chooselist[2]])
    header3.add_run(' - ')
    header3.add_run(cfg.bodylist[cfg.chooselist[2]])
    header3.runs[0].font.bold = True


#TODO Construct docx file


#print(cfg.chooselist)

def headerassemble():
    out = input('Pick a header: ')
    cfg.coverletter.add_paragraph(out)